import docx
from lxml import etree
import pandas as pd
import docx_path_finder
import os
from pathlib import Path

def extract_text_and_tables(docx_file, output_file_path):
    doc = docx.Document(docx_file)
    all_content_text = []
    
    for paragraph in doc.paragraphs:
        all_content_text.append((paragraph.text, paragraph._element))
    
    for i, table in enumerate(doc.tables):
        all_content_text.append((table, table._element))
    
    # Sort content based on the position in the document
    all_content_text.sort(key=lambda x: get_element_position(x[1]))
    
