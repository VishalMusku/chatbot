import docx
from lxml import etree
import docx_path_finder


def extract_text_and_tables(docx_file):
    doc = docx.Document(docx_file)
    all_content_text = []

    for paragraph in doc.paragraphs:
        all_content_text.append((paragraph.text, paragraph._element))

    for i,table in enumerate(doc.tables):
        all_content_text.append((table, table._element))

    # Sort content based on the position in the document
    all_content_text.sort(key=lambda x: get_element_position(x[1]))

    result_text = ""
    for item in all_content_text:
        if isinstance(item[0], str):
            result_text += item[0] + "\n"
        else:
            table_text = extract_table_text(item[0], i + 1)
            result_text += table_text + "\n"

    return result_text

def extract_table_text(table, table_number):
    table_text = f"Table {table_number}:\n"
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        row_text = "|".join(row_data)
        table_text += f"{row_text}\n"
    table_text += "\n" + "-" * 20 + "\n"
    return table_text

def get_element_position(element):
    """Get the position of the XML element in the document."""
    return int(element.xpath('count(./preceding::*)') + 1)


def main():
    
    docx_file_path = "{}".format(docx_path_finder.get_context()) # docx_path_finder.get_context()

    try:
        content_text = extract_text_and_tables(docx_file_path)
        
        # Optionally, you can save the combined text to a file
        with open('output_combined.txt', 'w') as file:
            file.write(content_text)
    except Exception as e:
        print(f"Error: {e}")
    


    